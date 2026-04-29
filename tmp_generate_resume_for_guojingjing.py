# -*- coding: utf-8 -*-
"""生成面向郭晶晶老师课题组的简历定制版。"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_PATH = Path(r"C:\Users\墨墨\Desktop\个人简历-郭晶晶课题组申请版.docx")


def set_font(run, name="宋体", size=11, bold=False, color=None):
    """统一字体与字号，确保导出的中文简历打开后不乱码。"""
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_doc_defaults(doc):
    """统一页边距与正文样式，影响范围限于简历排版。"""
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)


def add_title(doc, name, subtitle):
    """顶部标题区用于形成正式申请版简历的第一印象。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(name)
    set_font(r, name="黑体", size=20, bold=True, color=(20, 44, 86))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    set_font(r2, size=11, color=(80, 80, 80))


def add_contact_line(doc, text):
    """联系信息行用于压缩个人基础信息，影响范围限于页眉下方。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size=10)


def add_section(doc, title):
    """分节标题统一蓝色强调，便于老师快速浏览。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_font(r, name="黑体", size=13, bold=True, color=(20, 44, 86))


def add_paragraph(doc, text, first_indent=True):
    """正文段落承担经历和说明信息，影响范围限于内容区。"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.35
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_font(r, size=11)


def add_bullets(doc, items):
    """要点列表用于突出课程、技能和项目亮点。"""
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(item)
        set_font(r, size=11)


def add_two_col_table(doc, rows):
    """双列表格用于整齐展示课程成绩和获奖信息。"""
    table = doc.add_table(rows=0, cols=2)
    table.autofit = True
    for left, right in rows:
        row = table.add_row().cells
        row[0].width = Cm(5.5)
        row[1].width = Cm(10.5)
        row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p1 = row[0].paragraphs[0]
        p2 = row[1].paragraphs[0]
        p1.paragraph_format.line_spacing = 1.15
        p2.paragraph_format.line_spacing = 1.15
        r1 = p1.add_run(left)
        r2 = p2.add_run(right)
        set_font(r1, name="黑体", size=10, bold=True)
        set_font(r2, size=10)
    doc.add_paragraph("")


def build_resume():
    """生成针对郭晶晶老师课题组的申请版简历。"""
    doc = Document()
    set_doc_defaults(doc)

    add_title(doc, "金邦昱", "面向西安电子科技大学网信院郭晶晶老师课题组申请简历")
    add_contact_line(doc, "生日：2007.02.25    地址：西安电子科技大学    邮箱：jinbangyu@outlook.com")

    add_section(doc, "一、申请意向")
    add_paragraph(
        doc,
        "希望加入西安电子科技大学网络与信息安全学院郭晶晶老师课题组，重点关注信任与隐私、人工智能安全、Web 安全等方向。我对科研和项目实践有较强兴趣，希望在真实研究与项目训练中系统提升数学基础、编程能力、问题分析能力和安全方向的研究能力。",
    )

    add_section(doc, "二、教育背景")
    add_bullets(
        doc,
        [
            "西安电子科技大学，本科在读。",
            "目前重视数学、编程和英语基础训练，希望尽早进入科研与项目环境，在实践中学习更多知识。",
        ],
    )

    add_section(doc, "三、相关课程与成绩")
    add_two_col_table(
        doc,
        [
            ("高等数学（大一上期末）", "85"),
            ("线性代数", "92"),
            ("程序设计", "90"),
            ("大学英语", "89"),
            ("课程基础概括", "数学与编程基础较扎实，适合继续向网络安全、人工智能安全等方向深化。"),
        ],
    )

    add_section(doc, "四、技能基础")
    add_bullets(
        doc,
        [
            "编程语言：掌握 C 语言，能够完成基础程序设计与题目实现；掌握 Python 基本语法，能进行简单脚本编写与逻辑实现。",
            "英语能力：CET-4 597 分，具备较好的英文阅读能力，可继续加强对英文技术资料与论文的阅读。",
            "学习能力：愿意围绕安全方向继续补充数据结构、计算机网络、操作系统、Web 技术与网络空间安全基础知识。",
        ],
    )

    add_section(doc, "五、竞赛与获奖")
    add_bullets(
        doc,
        [
            "全国大学生英语竞赛三等奖。",
            "大学英语四级 CET-4：597 分。",
        ],
    )

    add_section(doc, "六、项目经历")
    add_paragraph(
        doc,
        "项目名称：XDU CampusMind 西电校园文创智能生成与展示平台",
        first_indent=False,
    )
    add_bullets(
        doc,
        [
            "围绕西安电子科技大学校园文化元素，参与完成元素图谱、AI 问答、文创创意生成与展示页面的项目实践。",
            "在项目过程中重点接触了需求梳理、页面内容组织、文案结构化表达、功能联调与结果优化等工作，对完整项目从想法到落地的过程有了更直观认识。",
            "通过该项目进一步提升了自己对前后端协作、提示词设计、交互逻辑与实际产品落地的理解，也增强了继续参与科研与项目训练的意愿。",
            "本项目在实现过程中由 Codex 辅助完成部分开发、调试与文档整理工作，我也在过程中主动学习和理解相关技术细节。"
        ],
    )

    add_section(doc, "七、与课题组方向的匹配点")
    add_bullets(
        doc,
        [
            "具备较好的数学、程序设计和英语基础，符合继续进入网络空间安全方向学习的前提条件。",
            "对信任与隐私、人工智能安全、Web 安全等方向有明确兴趣，希望在老师指导下逐步建立系统研究能力。",
            "有较强的项目实践热情，不满足于只停留在课堂学习，希望通过实际课题训练提升能力。",
            "愿意长期投入，先从基础做起，在阅读、实现、复现和小课题实践中逐步成长。"
        ],
    )

    add_section(doc, "八、自我评价")
    add_paragraph(
        doc,
        "我对科研与项目实践有较强热情，学习态度认真，愿意持续投入时间补齐基础、阅读资料、完成任务并主动请教。我希望进入一个要求明确、训练扎实的课题组，在真实研究和项目环境中不断学习更多知识，提升自己的综合能力。"
    )

    return doc


if __name__ == "__main__":
    build_resume().save(OUTPUT_PATH)
    print(OUTPUT_PATH)
