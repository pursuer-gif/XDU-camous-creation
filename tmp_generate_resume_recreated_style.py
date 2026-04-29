# -*- coding: utf-8 -*-
"""生成接近原简历框架感的课题组申请版简历。"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_PATH = Path(r"D:\vscodeprojects\campus\resume_recreated_style.docx")

NAVY = (23, 46, 89)
LIGHT = (242, 246, 252)
TEXT = (40, 40, 40)


def set_font(run, name="宋体", size=11, bold=False, color=TEXT):
    """统一中文字体和颜色，确保导出后中文显示稳定。"""
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(*color)


def set_border(cell, color="D5DCE8", size="8"):
    """单元格描边用于模拟原简历的框架线条感。"""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def shade_cell(cell, fill):
    """底色块用于重建原版简历中的装饰区域。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def clear_cell(cell):
    """清理默认内容，便于精确控制文本布局。"""
    cell.text = ""
    return cell.paragraphs[0]


def write_cell(cell, text, *, font="宋体", size=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=TEXT):
    """写入单元格文本，统一控制风格。"""
    p = clear_cell(cell)
    p.alignment = align
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_font(run, name=font, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_border(cell)


def add_para(cell, text, *, font="宋体", size=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=TEXT):
    """在单元格内追加段落，用于多行内容展示。"""
    p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_font(run, name=font, size=size, bold=bold, color=color)


def set_doc(doc):
    """页面设置整体偏紧凑，模拟原简历竖向排版节奏。"""
    section = doc.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)


def build():
    """重建带框架感的申请版简历，内容针对郭晶晶老师课题组。"""
    doc = Document()
    set_doc(doc)

    outer = doc.add_table(rows=1, cols=1)
    outer.autofit = True
    shell = outer.cell(0, 0)
    shade_cell(shell, "FFFFFF")
    set_border(shell, color="AEB8C8", size="14")
    shell.text = ""

    # 顶部主框架
    header = shell.add_table(rows=1, cols=2)
    header.autofit = False
    header.columns[0].width = Cm(4.2)
    header.columns[1].width = Cm(12.5)

    left = header.cell(0, 0)
    right = header.cell(0, 1)
    shade_cell(left, "EAF0FA")
    set_border(left, color="AEB8C8", size="10")
    set_border(right, color="AEB8C8", size="10")
    write_cell(left, "个人简历", font="黑体", size=19, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=NAVY)
    add_para(left, "Personal Resume", font="Calibri", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=NAVY)
    add_para(left, "课题组申请版", font="宋体", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=NAVY)

    write_cell(right, "金邦昱", font="黑体", size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=NAVY)
    add_para(
        right,
        "申请方向：西安电子科技大学网信院郭晶晶老师课题组",
        font="宋体",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=TEXT,
    )
    add_para(
        right,
        "生日：2007.02.25    地址：西安电子科技大学    邮箱：jinbangyu@outlook.com",
        size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=(90, 90, 90),
    )

    shell.add_paragraph("")

    # 主体双列布局
    body = shell.add_table(rows=1, cols=2)
    body.autofit = False
    body.columns[0].width = Cm(5.4)
    body.columns[1].width = Cm(11.3)
    left_col = body.cell(0, 0)
    right_col = body.cell(0, 1)
    left_col.text = ""
    right_col.text = ""
    set_border(left_col, color="AEB8C8", size="10")
    set_border(right_col, color="AEB8C8", size="10")
    shade_cell(left_col, "F7F9FD")

    # 左栏：信息卡
    def left_block(title, items):
        block = left_col.add_table(rows=1, cols=1)
        cell = block.cell(0, 0)
        shade_cell(cell, "F7F9FD")
        set_border(cell, color="C7D1E0", size="8")
        write_cell(cell, title, font="黑体", size=11, bold=True, color=NAVY)
        for item in items:
            add_para(cell, item, size=10.2)
        left_col.add_paragraph("")

    left_block("基础信息", [
        "学校：西安电子科技大学",
        "身份：本科生",
        "申请意向：本科阶段进入课题组",
        "目标方向：信任与隐私、人工智能安全、Web 安全",
    ])

    left_block("课程成绩", [
        "高等数学：85",
        "线性代数：92",
        "程序设计：90",
        "大学英语：89",
    ])

    left_block("技能与语言", [
        "掌握 C 语言",
        "掌握 Python 基本语法",
        "CET-4：597",
        "具备良好英文资料阅读基础",
    ])

    left_block("获奖情况", [
        "全国大学生英语竞赛三等奖",
        "英语四级成绩优秀",
    ])

    # 右栏：主叙事内容
    def right_section(title, paragraphs, bullets=None):
        card = right_col.add_table(rows=1, cols=1)
        cell = card.cell(0, 0)
        set_border(cell, color="C7D1E0", size="8")
        write_cell(cell, title, font="黑体", size=11.5, bold=True, color=NAVY)
        for text in paragraphs:
            add_para(cell, text, size=10.4)
        if bullets:
            for item in bullets:
                add_para(cell, f"• {item}", size=10.4)
        right_col.add_paragraph("")

    right_section(
        "申请意向",
        [
            "我希望加入西安电子科技大学网络与信息安全学院郭晶晶老师课题组，重点关注信任与隐私、人工智能安全、Web 安全等方向。对我而言，进入课题组不仅是一次学习机会，更是尽早进入科研和项目实践环境、系统提升自己能力的重要起点。",
        ],
    )

    right_section(
        "项目经历",
        [
            "项目名称：XDU CampusMind 西电校园文创智能生成与展示平台。",
            "该项目围绕西安电子科技大学校园文化元素展开，涵盖元素图谱、AI 问答、文创创意生成与展示等内容。通过参与这个项目，我对从需求梳理、页面表达、功能联调到结果优化的完整过程有了更具体的认识。",
        ],
        bullets=[
            "在项目过程中接触并理解了内容结构化表达、交互逻辑设计和功能联调的基本流程。",
            "通过项目实践增强了自己对真实产品开发和持续迭代的兴趣。",
            "本项目在实现过程中由 Codex 辅助完成部分开发、调试与文档整理工作，我也在过程中主动学习和理解相关技术细节。",
        ],
    )

    right_section(
        "与课题组方向的匹配",
        [],
        bullets=[
            "具备较好的数学、编程和英语基础，适合作为进入网络空间安全方向学习的起点。",
            "对信任与隐私、人工智能安全、Web 安全等方向有明确兴趣，希望在老师指导下逐步建立研究能力。",
            "对科研与项目实践有持续热情，愿意从基础做起，在阅读、实现、复现和训练中不断积累。",
        ],
    )

    right_section(
        "个人陈述",
        [
            "我对科研和项目实践始终保持较强兴趣，也愿意投入时间补齐基础、阅读资料、完成任务并主动请教。我希望进入一个要求明确、训练扎实的课题组，在实践中学习更多知识，在不断完成真实任务的过程中提升自己的能力。",
        ],
    )

    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("面向郭晶晶老师课题组的定制申请版简历")
    set_font(r, size=9.5, color=(110, 110, 110))

    return doc


if __name__ == "__main__":
    build().save(OUTPUT_PATH)
    print(OUTPUT_PATH)
