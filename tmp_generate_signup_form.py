# -*- coding: utf-8 -*-
"""生成西电校园文创大赛报名表，确保中文内容正常写入 Word。"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE_DIR = Path(r"D:\vscodeprojects\campus")
OUTPUT_PATH = BASE_DIR / "signup_form.docx"


def set_border(cell):
    """统一单元格边框样式，影响范围仅限报名表表格外观。"""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_font(run, name="宋体", size=12, bold=False):
    """统一中文字体设置，避免 Word 打开时字体回退导致显示异常。"""
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def clear_cell(cell):
    """清空单元格默认段落，便于按模板重写内容。"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def write_cell(cell, text, *, align=WD_ALIGN_PARAGRAPH.CENTER, bold=False, font="宋体", size=11):
    """写入单元格文本并统一排版，影响范围限于报名表内容。"""
    p = clear_cell(cell)
    p.alignment = align
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_font(run, name=font, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_border(cell)


def append_paragraph(cell, text, *, align=WD_ALIGN_PARAGRAPH.LEFT, font="宋体", size=11):
    """在单元格内追加段落，便于多行选项和说明展示。"""
    p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_font(run, name=font, size=size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_border(cell)


def set_page(doc):
    """页面设置按常规 A4 中文文档处理，影响范围限于导出版式。"""
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)


def build_document():
    """生成报名表正文，作品信息按当前项目预填，个人信息保留可编辑空位。"""
    doc = Document()
    set_page(doc)

    p_attach = doc.add_paragraph()
    run_attach = p_attach.add_run("附件")
    set_font(run_attach, name="宋体", size=12)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("西安电子科技大学 2026 校园文创大赛报名表")
    set_font(run_title, name="黑体", size=16, bold=True)

    table = doc.add_table(rows=9, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(3.6), Cm(2.6), Cm(4.1), Cm(2.6), Cm(5.9)]
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = width

    # Row 1
    write_cell(table.cell(0, 0), "作品名称", bold=True, font="黑体", size=12)
    merged = table.cell(0, 1).merge(table.cell(0, 4))
    write_cell(
        merged,
        "XDU CampusMind 西电校园文创智能生成与展示平台",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        size=11,
    )

    # Row 2
    write_cell(table.cell(1, 0), "设计主题", bold=True, font="黑体", size=12)
    merged = table.cell(1, 1).merge(table.cell(1, 4))
    write_cell(merged, "☑ 红色基因   ☑ 科技特色   ☑ 文化传承", align=WD_ALIGN_PARAGRAPH.LEFT, size=11)
    append_paragraph(merged, "☑ 校园生活   ☐ 未来畅想", align=WD_ALIGN_PARAGRAPH.LEFT, size=11)

    # Row 3
    write_cell(table.cell(2, 0), "作品类别", bold=True, font="黑体", size=12)
    merged = table.cell(2, 1).merge(table.cell(2, 4))
    write_cell(
        merged,
        "☐ 服饰配件类   ☐ 创意礼品类   ☐ 办公用品类",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        size=11,
    )
    append_paragraph(
        merged,
        "☐ 艺术收藏类   ☐ 智能科技类   ☑ 数字内容类",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        size=11,
    )

    # Row 4
    write_cell(table.cell(3, 0), "是否申报 AI 设计赛道", bold=True, font="黑体", size=12)
    merged = table.cell(3, 1).merge(table.cell(3, 4))
    write_cell(
        merged,
        "☑ 是（请填写所用 AI 工具/技术：GPT-5.4、Gemini、百度千帆、React、Node.js）",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        size=10,
    )
    append_paragraph(merged, "☐ 否", align=WD_ALIGN_PARAGRAPH.LEFT, size=11)

    # Row 5
    write_cell(table.cell(4, 0), "作品简介", bold=True, font="黑体", size=12)
    merged = table.cell(4, 1).merge(table.cell(4, 4))
    write_cell(
        merged,
        "围绕西电校园文化元素构建集元素图谱、AI问答、文创创意生成与展示于一体的数字文创平台。（限 100 字）",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        size=11,
    )
    merged.height = Cm(4.0)

    # Participant block
    left = table.cell(5, 0).merge(table.cell(7, 0))
    write_cell(left, "参赛者信息", bold=True, font="黑体", size=12)

    write_cell(table.cell(5, 1), "姓名", bold=True, font="黑体", size=12)
    write_cell(table.cell(5, 2), "（团队填写团队名称）", size=10)
    write_cell(table.cell(5, 3), "身份", bold=True, font="黑体", size=12)
    write_cell(table.cell(5, 4), "教师 / 学生 / 校友 / 社会设计爱好者", align=WD_ALIGN_PARAGRAPH.LEFT, size=10)

    write_cell(table.cell(6, 1), "院系、年级、专业", bold=True, font="黑体", size=12)
    merged = table.cell(6, 2).merge(table.cell(6, 4))
    write_cell(
        merged,
        "（仅师生校友填写，教师填写所在单位，团队填写队长信息）",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        size=10,
    )

    write_cell(table.cell(7, 1), "联系电话", bold=True, font="黑体", size=12)
    write_cell(table.cell(7, 2), "", size=11)
    write_cell(table.cell(7, 3), "邮箱", bold=True, font="黑体", size=12)
    write_cell(table.cell(7, 4), "", size=11)

    # Row 8
    write_cell(table.cell(8, 0), "指导教师", bold=True, font="黑体", size=12)
    merged = table.cell(8, 1).merge(table.cell(8, 2))
    write_cell(merged, "", size=11)
    write_cell(table.cell(8, 3), "指导教师联系电话", bold=True, font="黑体", size=12)
    write_cell(table.cell(8, 4), "", size=11)

    # Originality block
    table2 = doc.add_table(rows=2, cols=5)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.autofit = False
    for row in table2.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = width

    write_cell(table2.cell(0, 0), "原创性声明", bold=True, font="黑体", size=12)
    merged = table2.cell(0, 1).merge(table2.cell(0, 4))
    write_cell(
        merged,
        "本人（团队）承诺参赛作品为原创设计，未侵犯任何第三方知识产权或其他合法权益，并同意主办方对作品进行宣传、生产及衍生开发。若使用 AI 工具，已确保内容符合版权规范。",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        size=11,
    )

    merged = table2.cell(1, 0).merge(table2.cell(1, 4))
    write_cell(merged, "签名：____________________            日期：____________________", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

    return doc


if __name__ == "__main__":
    build_document().save(OUTPUT_PATH)
    print(OUTPUT_PATH)
