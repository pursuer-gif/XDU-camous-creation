# -*- coding: utf-8 -*-
"""生成更适合研究组申请的简洁学术风简历。"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_PATH = Path(r"C:\Users\墨墨\Desktop\个人简历-郭晶晶课题组-简洁版.docx")


def set_font(run, name="宋体", size=11, bold=False, color=None):
    """统一字体设置，确保中文显示稳定。"""
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_doc_defaults(doc):
    """统一页面边距和正文样式，影响范围限于新简历版式。"""
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)


def set_cell_border(cell, color="D0D7E2"):
    """轻边框仅用于增强层次，不使用重装饰元素。"""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def shade_cell(cell, fill="F3F6FB"):
    """浅色底用于标题栏，保持简洁学术风。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def write_cell(cell, text, *, font="宋体", size=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    """统一写入单元格文本，保持简历整体排版一致。"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_font(run, name=font, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_border(cell)


def add_header(doc):
    """顶部信息区用于建立简洁、可信的第一印象。"""
    table = doc.add_table(rows=2, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("金邦昱")
    set_font(r, name="黑体", size=20, bold=True, color=(20, 44, 86))
    set_cell_border(cell, "FFFFFF")

    cell2 = table.cell(1, 0)
    cell2.text = ""
    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run("申请方向：西安电子科技大学网信院郭晶晶老师课题组 | 邮箱：jinbangyu@outlook.com | 地址：西安电子科技大学")
    set_font(r2, size=10.5, color=(90, 90, 90))
    set_cell_border(cell2, "FFFFFF")
    doc.add_paragraph("")


def add_section_title(doc, title):
    """分节标题采用浅底栏，提升信息识别度。"""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade_cell(cell)
    write_cell(cell, title, font="黑体", size=11.5, bold=True, color=(20, 44, 86))


def add_body_paragraph(doc, text):
    """正文段落用于承接申请动机与自我评价。"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_font(run, size=10.5)


def add_bullets(doc, items):
    """简洁项目符号列表用于突出事实信息。"""
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(item)
        set_font(run, size=10.5)


def add_two_col_table(doc, rows):
    """双列表格用于课程、技能等清晰表达。"""
    table = doc.add_table(rows=0, cols=2)
    table.autofit = True
    for left, right in rows:
        cells = table.add_row().cells
        write_cell(cells[0], left, font="黑体", size=10.3, bold=True)
        write_cell(cells[1], right, size=10.3)
    doc.add_paragraph("")


def build_resume():
    """生成面向课题组联系的清爽学术风简历。"""
    doc = Document()
    set_doc_defaults(doc)
    add_header(doc)

    add_section_title(doc, "申请意向")
    add_body_paragraph(
        doc,
        "希望加入西安电子科技大学网络与信息安全学院郭晶晶老师课题组，重点关注信任与隐私、人工智能安全、Web 安全等方向。我对科研与项目实践有较强热情，希望在真实研究和项目训练中不断学习更多知识，系统提升数学基础、编程能力与安全方向研究能力。",
    )

    add_section_title(doc, "教育背景")
    add_bullets(
        doc,
        [
            "西安电子科技大学，本科在读。",
            "重视数学、编程与英语基础训练，希望尽早进入科研环境开展实践学习。",
        ],
    )

    add_section_title(doc, "课程基础")
    add_two_col_table(
        doc,
        [
            ("高等数学（大一上期末）", "85"),
            ("线性代数", "92"),
            ("程序设计", "90"),
            ("大学英语", "89"),
        ],
    )

    add_section_title(doc, "技能与语言能力")
    add_bullets(
        doc,
        [
            "掌握 C 语言，能够完成基础程序设计与简单算法实现。",
            "掌握 Python 基本语法，能够进行基础脚本编写和逻辑实现。",
            "英语能力较好，CET-4 597 分，具备较好的英文技术资料阅读基础。",
        ],
    )

    add_section_title(doc, "获奖与成绩")
    add_bullets(
        doc,
        [
            "全国大学生英语竞赛三等奖。",
            "大学英语四级（CET-4）：597 分。",
        ],
    )

    add_section_title(doc, "项目经历")
    add_bullets(
        doc,
        [
            "项目名称：XDU CampusMind 西电校园文创智能生成与展示平台。",
            "围绕西安电子科技大学校园文化元素，参与完成元素图谱、AI 问答、文创创意生成与展示页面的项目实践。",
            "在项目中接触了需求梳理、页面内容组织、交互逻辑、文案结构化表达与功能联调，对完整项目从构思到落地有了更具体的认识。",
            "项目在实现过程中由 Codex 辅助完成部分开发、调试与文档整理工作，我也在过程中持续学习并理解相关技术细节。"
        ],
    )

    add_section_title(doc, "与课题组方向的匹配")
    add_bullets(
        doc,
        [
            "数学、编程和英语基础较扎实，适合作为进入网络空间安全方向学习的起点。",
            "对信任与隐私、人工智能安全、Web 安全等方向有明确兴趣，希望在老师指导下系统学习。",
            "对科研与项目实践有持续热情，愿意从基础做起，在阅读、实现、复现与小课题训练中不断进步。",
        ],
    )

    add_section_title(doc, "个人陈述")
    add_body_paragraph(
        doc,
        "我对科研与项目实践始终保持较强兴趣，学习态度认真，也愿意为长期积累投入时间。我希望进入一个要求明确、训练扎实的课题组，在实践中学习更多知识，在项目和研究中不断提升自己。"
    )

    return doc


if __name__ == "__main__":
    build_resume().save(OUTPUT_PATH)
    print(OUTPUT_PATH)
