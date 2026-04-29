# -*- coding: utf-8 -*-
"""重新生成比赛文档，确保中文内容以 UTF-8 正常写入 Word。"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE_DIR = Path(r"D:\vscodeprojects\campus")


def set_cell_shading(cell, fill):
    """表头底色仅用于增强可读性，影响范围限于文档表格样式。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_doc_defaults(doc):
    """统一页边距和中文字体，确保导出的 Word 在本地打开时显示稳定。"""
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)


def add_title(doc, title, subtitle=None):
    """标题区用于形成正式比赛材料版式，影响范围限于首页抬头。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(18)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.name = "宋体"
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r2.font.size = Pt(11)


def add_heading(doc, text, level=1):
    """章节标题统一使用黑体，便于比赛评审快速扫描结构。"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(15 if level == 1 else 13)


def add_paragraph(doc, text, first_line_cm=0.74):
    """正文统一首行缩进与行距，保证正式文档阅读体验。"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(first_line_cm)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(12)


def add_bullets(doc, items):
    """要点列表用于压缩说明信息，影响范围限于说明条目展示。"""
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.35
        r = p.add_run(item)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(12)


def add_table(doc, rows):
    """参数表用于技术文档的结构化说明，影响范围限于表格区域。"""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "项目"
    hdr[1].text = "说明"
    for cell in hdr:
        set_cell_shading(cell, "D9EAF7")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "黑体"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
                r.font.size = Pt(11)
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "宋体"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    r.font.size = Pt(11)
    doc.add_paragraph("")


def build_design_doc():
    """设计说明文档用于覆盖赛事要求里的设计说明材料。"""
    doc = Document()
    set_doc_defaults(doc)
    add_title(doc, "西安电子科技大学 2026 校园文创大赛", "设计说明文档")
    add_paragraph(
        doc,
        "项目名称为“XDU CampusMind 西电校园文创智能生成与展示平台”。本方案围绕西安电子科技大学校园文化元素展开，以“半部电台起家、银杏大道、图书馆、通信天线、集成电路、厚德求真”等核心符号为内容基础，构建集校园元素图谱、AI问答、文创创意生成与展示于一体的数字化文创展示产品。",
    )
    add_heading(doc, "一、设计主题对应说明")
    add_bullets(
        doc,
        [
            "红色基因：以“半部电台起家”为精神起点，突出西电从革命通信传统中走来的校史脉络。",
            "科技特色：围绕电子信息、通信报国、集成电路等学科特色，形成高识别度的校园科技表达。",
            "文化传承：将校史叙事、校园空间、校训精神与现代数字交互方式结合，增强传播力。",
            "校园生活：通过银杏大道、图书馆等具象场景补充校园温度，避免表达过于抽象。",
            "未来构想：通过智能问答和创意生成能力，将传统校园文化延展到未来数字文创场景。",
        ],
    )
    add_heading(doc, "二、作品定位")
    add_paragraph(
        doc,
        "本作品属于数字内容类与AI设计赛道结合的校园文创作品。产品以网页端交互展示为主体，兼顾文创方案生成、视觉方案解释与互动问答能力，适合作为校园文化展示、文创提案展示及后续成果延展的数字底座。",
    )
    add_heading(doc, "三、核心创意说明")
    add_paragraph(
        doc,
        "本设计不是简单罗列校园元素，而是先把西电的文化内容拆解成“精神源点、校园记忆、学科语言、空间场景、价值表达”五个层次，再通过统一视觉语言和智能交互方式重新组织，让用户能同时理解“西电是谁、为什么这样表达、这些元素如何转化为文创设计”。",
    )
    add_bullets(
        doc,
        [
            "元素图谱：将校史与校园文化符号结构化整理，形成可浏览、可解释的文化地图。",
            "西小电问答：点击校园智能助理形象后，可围绕校史、元素含义、视觉转译与页面表达进行自然语言问答。",
            "创意生成：输入主题、风格、载体和校园元素后，可生成设计概念、视觉元素、应用场景、提示词与图像预览。",
            "成果展示：保留历史记录与精选画廊，使网页本身具备展示板和提案库的双重作用。",
        ],
    )
    add_heading(doc, "四、视觉与交互设计说明")
    add_bullets(
        doc,
        [
            "视觉基调采用科技感、秩序感与校园记忆并置的方式，以深空蓝、信号青、西电蓝为主，辅以银杏金。",
            "页面结构遵循“先认识西电，再理解元素，再进入生成，再沉淀结果”的自然路径。",
            "右下角悬浮的“西小电”承担问答入口，降低首次使用门槛，增强内容解释力。",
            "设计避免泛校园化，强调电波、线路、网格、建筑轮廓等具有西电辨识度的视觉语汇。",
        ],
    )
    add_heading(doc, "五、AI设计赛道说明")
    add_paragraph(
        doc,
        "本作品符合赛事对AI创新应用的要求。AI技术并非作为装饰性附加，而是实际参与了校园文化信息整理、交互问答、创意提示生成、图像方案生成和成果展示等关键流程。作品中明确保留了AI参与链路、生成结果和应用边界，确保技术使用清晰、合理、可说明。",
    )
    add_heading(doc, "六、原创性与合规说明")
    add_bullets(
        doc,
        [
            "作品围绕西安电子科技大学校园文化语境自主组织内容与交互结构，不是套用通用校园模板。",
            "涉及校史与学校形象时，采用高置信信息与官方视觉规范约束，避免错误延展。",
            "涉及校徽或学校官方视觉识别时，严格遵循学校VIS要求，不自行臆造或变形。",
            "AI生成内容仅作为创意辅助与表达增强工具，最终页面结构、交互链路和产品整合经过人工判断与修改完成。",
        ],
    )
    add_heading(doc, "七、开发协作说明")
    add_paragraph(
        doc,
        "本项目在开发协作上采用多模型协同方式完成：GPT-5.4 主要负责提示词设计与生成链路中的文案约束梳理，Gemini 参与前端界面实现与部分页面交互组织，我负责后端接口、聊天服务、图片生成链路、前后端联调及整体稳定性修正。最终产品的可运行版本、后端能力落地、联网问答稳定性与关键问题修复，主要由我完成修改与整合。",
    )
    add_heading(doc, "八、交付内容说明")
    add_bullets(
        doc,
        [
            "网页交互成品：用于展示校园文创生成与问答功能。",
            "设计说明文档：用于阐明作品主题、结构、创意来源与AI使用方式。",
            "技术方案文档：用于说明系统架构、模型接入、接口设计与实现方式。",
            "可扩展的图像预览结果与历史记录：用于后续打样、答辩展示和成果转化。",
        ],
    )
    add_heading(doc, "九、总结")
    add_paragraph(
        doc,
        "本方案以“文化内容结构化 + AI交互化 + 文创表达可视化”为核心，把西电校园文化从静态素材转化为可解释、可交互、可生成、可展示的数字文创产品，既符合赛事导向，也具备后续展示与延展价值。",
    )
    return doc


def build_tech_doc():
    """技术方案文档用于覆盖赛事要求里的技术实现说明材料。"""
    doc = Document()
    set_doc_defaults(doc)
    add_title(doc, "西安电子科技大学 2026 校园文创大赛", "技术方案文档")
    add_paragraph(
        doc,
        "本技术方案对应“XDU CampusMind 西电校园文创智能生成与展示平台”。系统采用前后端分离结构，前端负责校园元素展示、方案交互和结果呈现，后端负责生成接口、聊天接口、图片生成能力和历史数据管理。技术目标是在可展示、可交互、可扩展的前提下，完成校园文创生成与智能问答的整体交付。",
    )
    add_heading(doc, "一、总体技术目标")
    add_bullets(
        doc,
        [
            "构建一个面向校园文创场景的网页产品，支持文化展示、创意生成与AI问答。",
            "提供稳定的“元素图谱 + 西小电问答 + 方案生成 + 图像预览 + 历史沉淀”闭环。",
            "保证前端演示链路清晰、后端调用稳定、关键接口具备限流和异常兜底能力。",
        ],
    )
    add_heading(doc, "二、系统架构")
    add_table(
        doc,
        [
            ("前端层", "基于 React + TypeScript + Vite 构建，负责页面展示、输入交互、聊天浮窗、结果渲染和本地静态预览服务。"),
            ("后端层", "基于 Node.js + Express 构建，负责 /api/chat、/api/generate、/api/history、/api/featured 等接口。"),
            ("AI服务层", "接入百度千帆兼容接口，承担聊天问答与图片生成能力。"),
            ("数据层", "当前以项目内历史记录与精选结果存储为主，满足展示与回看需求。"),
            ("代理层", "本地静态服务通过同源 /api 代理转发到 127.0.0.1:3000，减少前端跨端口失败。"),
        ],
    )
    add_heading(doc, "三、核心模块设计")
    add_bullets(
        doc,
        [
            "元素图谱模块：展示西电校园文化元素、时间线和转译路径，承担文化解释入口。",
            "西小电聊天模块：前端通过 /api/chat 调用后端聊天服务，支持多轮消息、失败重试和本地兜底。",
            "方案生成模块：根据主题、风格、载体和校园元素生成设计概念、视觉要素和图像提示词。",
            "图片生成模块：调用千帆生图接口生成主图与多方向预览图。",
            "历史与精选模块：保存生成记录和精选展示结果，用于回看与展示。",
        ],
    )
    add_heading(doc, "四、聊天服务实现方案")
    add_bullets(
        doc,
        [
            "聊天服务通过后端统一封装，避免前端直接暴露模型调用细节。",
            "后端对消息进行清洗，仅保留 user/assistant 角色，并控制长度，确保输入稳定。",
            "请求失败时，后端先进行同模型重试，再切换备用模型，最后才回退到本地讲解模式。",
            "当前默认主模型为 ernie-4.5-turbo-32k，备用模型为 ernie-x1-turbo-32k，用于降低偶发回退概率。",
            "前端对聊天请求也保留一次短重试，并在本次请求失败时明确提示，避免状态误导。",
        ],
    )
    add_heading(doc, "五、图片生成方案")
    add_bullets(
        doc,
        [
            "图片生成使用千帆图片生成接口，根据产品类型自动匹配尺寸。",
            "系统会先生成主图，再逐个生成多方向预览图，避免一次并发造成更高失败率。",
            "当前图片生成已经可用，但仍受第三方服务限流影响，属于可展示级稳定，不是强生产级稳定。",
        ],
    )
    add_heading(doc, "六、前后端协作说明")
    add_paragraph(
        doc,
        "在协作分工上，GPT-5.4 主要参与提示词设计、生成约束整理与问答角色设定；Gemini 参与前端页面结构与交互界面实现；我负责后端接口开发、千帆聊天接入、图片生成链路、前端本地代理修正、异常处理、模型回退、接口联调与最终稳定性修复。最终产品能够正常运行并达到当前展示效果，主要由我完成关键修改与整合。",
    )
    add_heading(doc, "七、环境与依赖说明")
    add_table(
        doc,
        [
            ("运行环境", "Node.js >= 18"),
            ("前端框架", "React 18、TypeScript、Vite、Tailwind CSS"),
            ("后端框架", "Node.js、Express、undici"),
            ("第三方能力", "百度千帆聊天接口、百度千帆图片生成接口"),
            ("关键环境变量", "QIANFAN_API_KEY、QIANFAN_MODEL、QIANFAN_FALLBACK_MODEL、QIANFAN_IMAGE_MODEL"),
        ],
    )
    add_heading(doc, "八、当前实现状态说明")
    add_bullets(
        doc,
        [
            "西小电聊天功能已联通，可通过页面右下角入口进行联网问答。",
            "本地静态服务已通过同源代理转发后端接口，减少浏览器端跨端口调用失败。",
            "方案生成接口可正常输出结构化结果，并能附带图片生成结果。",
            "生图能力已恢复可用，但在高频调用下仍可能受到第三方服务限流影响。",
            "整体状态适合作为答辩展示、成果演示和终稿提交材料。",
        ],
    )
    add_heading(doc, "九、风险与应对")
    add_bullets(
        doc,
        [
            "第三方模型偶发波动：通过后端重试、备用模型切换和本地兜底减小影响。",
            "图片生成限流：采用逐张生成策略，避免一次性并发过高。",
            "本机 localhost 路径不稳定：本地代理默认改为 127.0.0.1，降低请求失败概率。",
            "页面状态误导：前端已改为只提示“本次请求”结果，避免旧失败状态残留。",
        ],
    )
    add_heading(doc, "十、总结")
    add_paragraph(
        doc,
        "本技术方案以稳定交付为核心，优先保证聊天链路、页面展示链路和生成链路的可运行性。当前版本已经满足比赛展示、终稿演示和成果说明需求，并为后续继续扩展真实大模型生成能力保留了接口与结构基础。",
    )
    return doc


if __name__ == "__main__":
    build_design_doc().save(BASE_DIR / "design_description.docx")
    build_tech_doc().save(BASE_DIR / "technical_solution.docx")
    print("done")
